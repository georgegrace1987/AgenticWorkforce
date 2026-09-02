from io import BytesIO

from docx import Document


def extract_docx_text(content: bytes) -> str:
    """Extract paragraphs and table rows while preserving document order enough for analysis."""
    document = Document(BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)