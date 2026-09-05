from pathlib import Path
from typing import Iterable, cast

from docx.api import Document as create_document
from docx.document import Document
from docx.styles.style import ParagraphStyle
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.models.requirements import RequirementState


class SRDGenerator:
    """
    Generate an enterprise-grade Software Requirements Specification.

    The generator is intentionally defensive:
    - Never leaves meaningful document sections silently empty.
    - Uses TBD when information is unavailable.
    - Derives lightweight metadata from existing requirement text.
    - Keeps the document suitable for downstream test-case generation.
    """

    TBD = "TBD"

    def generate(self, state: RequirementState, output_path: Path) -> Path:
        document: Document = create_document()

        self._configure_page(document)
        self._configure_styles(document)

        self._add_cover(document, state)
        self._add_document_information(document, state)

        # ---------------------------------------------------------
        # 1. Executive Summary
        # ---------------------------------------------------------
        self._heading(document, "1. Executive Summary")

        self._text(
            document,
            state.business_objective,
            fallback=(
                "The objective of this solution has not yet been fully "
                "defined. Refer to the captured requirements and open "
                "questions for the current scope."
            ),
        )

        # ---------------------------------------------------------
        # 2. Scope
        # ---------------------------------------------------------
        self._heading(document, "2. Scope")

        self._subheading(document, "2.1 In Scope")

        self._list(
            document,
            state.functional_requirements,
            fallback="Functional scope has not yet been explicitly defined.",
        )

        self._subheading(document, "2.2 Out of Scope")

        self._paragraph(
            document,
            "The following items are explicitly outside the currently "
            "identified scope:"
        )

        self._bullet(
            document,
            "No explicit out-of-scope items have been identified.",
        )

        # ---------------------------------------------------------
        # 3. Stakeholders and User Roles
        # ---------------------------------------------------------
        self._heading(document, "3. Stakeholders and User Roles")

        stakeholder_rows = []

        for stakeholder in self._clean(state.stakeholders):
            stakeholder_rows.append(
                [
                    "Stakeholder",
                    stakeholder,
                    self._stakeholder_responsibility(stakeholder),
                ]
            )

        for role in self._clean(state.user_roles):
            stakeholder_rows.append(
                [
                    "User Role",
                    role,
                    self._role_responsibility(role),
                ]
            )

        if not stakeholder_rows:
            stakeholder_rows = [
                [
                    self.TBD,
                    self.TBD,
                    "Stakeholder and user responsibilities require clarification.",
                ]
            ]

        self._table(
            document,
            ["Type", "Name or Role", "Responsibilities"],
            stakeholder_rows,
        )

        # ---------------------------------------------------------
        # 4. Functional Requirements
        # ---------------------------------------------------------
        self._heading(document, "4. Functional Requirements")

        self._add_functional_requirements(
            document,
            state.functional_requirements,
        )

        # ---------------------------------------------------------
        # 5. Non-Functional Requirements
        # ---------------------------------------------------------
        self._heading(document, "5. Non-Functional Requirements")

        self._add_non_functional_requirements(
            document,
            state.non_functional_requirements,
        )

        # ---------------------------------------------------------
        # 6. Technical Requirements
        # ---------------------------------------------------------
        self._heading(document, "6. Technical Requirements")

        self._list(
            document,
            state.technical_requirements,
            fallback="Technical requirements have not yet been defined.",
        )

        # ---------------------------------------------------------
        # 7. Technical Dependencies
        # ---------------------------------------------------------
        self._heading(document, "7. Technical Dependencies")

        self._list(
            document,
            state.dependencies,
            fallback="No technical dependencies have been explicitly identified.",
        )

        # ---------------------------------------------------------
        # 8. Exceptions and Error Handling
        # ---------------------------------------------------------
        self._heading(document, "8. Exceptions and Error Handling")

        self._list(
            document,
            state.exceptions,
            fallback=(
                "Specific exception and error-handling requirements "
                "have not yet been defined."
            ),
        )

        # ---------------------------------------------------------
        # 9. Restrictions and Constraints
        # ---------------------------------------------------------
        self._heading(document, "9. Restrictions and Constraints")

        restrictions = (
            self._clean(state.restrictions)
            + self._clean(state.constraints)
        )

        self._list(
            document,
            restrictions,
            fallback="No explicit restrictions or constraints have been identified.",
        )

        # ---------------------------------------------------------
        # 10. Assumptions
        # ---------------------------------------------------------
        self._heading(document, "10. Assumptions")

        self._list(
            document,
            state.assumptions,
            fallback="No explicit assumptions have been recorded.",
        )

        # ---------------------------------------------------------
        # 11. Risks
        # ---------------------------------------------------------
        self._heading(document, "11. Risks")

        self._list(
            document,
            state.risks,
            fallback="No project or technical risks have been explicitly identified.",
        )

        # ---------------------------------------------------------
        # 12. Traceability and Acceptance
        # ---------------------------------------------------------
        self._heading(document, "12. Traceability and Acceptance")

        self._add_traceability(
            document,
            state,
        )

        # ---------------------------------------------------------
        # 13. Open Issues
        # ---------------------------------------------------------
        self._heading(document, "13. Open Issues")

        self._list(
            document,
            state.open_questions,
            fallback="No open questions have been recorded.",
        )

        # ---------------------------------------------------------
        # 14. Requirement Quality Summary
        # ---------------------------------------------------------
        self._heading(document, "14. Requirement Quality Summary")

        self._add_quality_summary(
            document,
            state,
        )

        # ---------------------------------------------------------
        # 15. Test Automation Readiness
        # ---------------------------------------------------------
        self._heading(document, "15. Test Automation Readiness")

        self._add_test_readiness(
            document,
            state,
        )

        # ---------------------------------------------------------
        # Footer / metadata
        # ---------------------------------------------------------
        self._configure_footer(document)

        document.core_properties.title = (
            "Software Requirements Specification"
        )

        document.core_properties.subject = (
            "Enterprise software requirements"
        )

        document.core_properties.author = (
            "AI Requirements Wizard"
        )

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(str(output_path))

        return output_path

    # =============================================================
    # DOCUMENT SETUP
    # =============================================================

    @staticmethod
    def _configure_page(document: Document) -> None:
        section = document.sections[0]

        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    @staticmethod
    def _configure_styles(document: Document) -> None:
        normal = cast(ParagraphStyle, document.styles["Normal"])
        normal.font.name = "Aptos"
        normal.font.size = Pt(10)

        for name, size in (
            ("Title", 25),
            ("Heading 1", 16),
            ("Heading 2", 12),
        ):
            style = cast(ParagraphStyle, document.styles[name])
            style.font.name = "Aptos Display"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(31, 78, 95)

    # =============================================================
    # COVER
    # =============================================================

    def _add_cover(
        self,
        document: Document,
        state: RequirementState,
    ) -> None:
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title.add_run(
            "Software Requirements Specification"
        )

        subtitle = document.add_paragraph()

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = subtitle.add_run(
            "AI Requirements Wizard"
        )

        run.italic = True

        objective = (
            state.business_objective
            or "Requirements definition document"
        )

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.add_run(objective)

        document.add_paragraph()

    def _add_document_information(
        self,
        document: Document,
        state: RequirementState,
    ) -> None:
        self._table(
            document,
            ["Document Information", "Value"],
            [
                [
                    "Project Objective",
                    self._value(state.business_objective),
                ],
                [
                    "Document Status",
                    "Draft",
                ],
                [
                    "Version",
                    "1.0",
                ],
                [
                    "Prepared By",
                    "AI Requirements Wizard",
                ],
                [
                    "Functional Requirements",
                    str(len(self._clean(
                        state.functional_requirements
                    ))),
                ],
                [
                    "Non-Functional Requirements",
                    str(len(self._clean(
                        state.non_functional_requirements
                    ))),
                ],
                [
                    "Open Questions",
                    str(len(self._clean(
                        state.open_questions
                    ))),
                ],
            ],
        )

    # =============================================================
    # REQUIREMENTS
    # =============================================================

    def _add_functional_requirements(
        self,
        document: Document,
        requirements: list[str],
    ) -> None:
        values = self._clean(requirements)

        if not values:
            self._paragraph(
                document,
                "No functional requirements have been captured.",
            )
            return

        rows = []

        for index, requirement in enumerate(values, 1):
            priority = self._infer_priority(requirement)

            acceptance = self._generate_acceptance_criteria(
                requirement
            )

            rows.append(
                [
                    f"FR-{index:03d}",
                    requirement,
                    priority,
                    acceptance,
                ]
            )

        self._table(
            document,
            [
                "ID",
                "Requirement",
                "Priority",
                "Acceptance Criteria",
            ],
            rows,
        )

    def _add_non_functional_requirements(
        self,
        document: Document,
        requirements: list[str],
    ) -> None:
        values = self._clean(requirements)

        if not values:
            self._paragraph(
                document,
                "No non-functional requirements have been captured.",
            )
            return

        rows = []

        for index, requirement in enumerate(values, 1):
            category = self._infer_nfr_category(
                requirement
            )

            target = self._infer_nfr_target(
                requirement,
                category,
            )

            rows.append(
                [
                    f"NFR-{index:03d}",
                    category,
                    requirement,
                    target,
                ]
            )

        self._table(
            document,
            [
                "ID",
                "Category",
                "Requirement",
                "Measure or Target",
            ],
            rows,
        )

    # =============================================================
    # TRACEABILITY
    # =============================================================

    def _add_traceability(
        self,
        document: Document,
        state: RequirementState,
    ) -> None:
        rows = []

        functional = self._clean(
            state.functional_requirements
        )

        non_functional = self._clean(
            state.non_functional_requirements
        )

        for index, requirement in enumerate(
            functional,
            1,
        ):
            rows.append(
                [
                    f"FR-{index:03d}",
                    requirement,
                    self._generate_acceptance_criteria(
                        requirement
                    ),
                    self._verification_method(
                        requirement
                    ),
                ]
            )

        for index, requirement in enumerate(
            non_functional,
            1,
        ):
            rows.append(
                [
                    f"NFR-{index:03d}",
                    requirement,
                    self._generate_acceptance_criteria(
                        requirement
                    ),
                    self._verification_method(
                        requirement
                    ),
                ]
            )

        if not rows:
            rows.append(
                [
                    self.TBD,
                    "No requirements captured.",
                    self.TBD,
                    self.TBD,
                ]
            )

        self._table(
            document,
            [
                "Requirement ID",
                "Requirement",
                "Acceptance Criteria",
                "Verification Method",
            ],
            rows,
        )

    # =============================================================
    # QUALITY
    # =============================================================

    def _add_quality_summary(
        self,
        document: Document,
        state: RequirementState,
    ) -> None:
        functional = self._clean(
            state.functional_requirements
        )

        non_functional = self._clean(
            state.non_functional_requirements
        )

        open_questions = self._clean(
            state.open_questions
        )

        total = (
            len(functional)
            + len(non_functional)
        )

        quality_rows = [
            [
                "Total Requirements",
                str(total),
            ],
            [
                "Functional Requirements",
                str(len(functional)),
            ],
            [
                "Non-Functional Requirements",
                str(len(non_functional)),
            ],
            [
                "Open Questions",
                str(len(open_questions)),
            ],
            [
                "Requirements With Explicit Priority",
                str(
                    sum(
                        self._has_explicit_priority(
                            item
                        )
                        for item in functional
                    )
                ),
            ],
        ]

        self._table(
            document,
            ["Metric", "Value"],
            quality_rows,
        )

        if open_questions:
            self._paragraph(
                document,
                (
                    "The document should not be considered "
                    "implementation-ready until the open questions "
                    "have been resolved."
                ),
            )

    # =============================================================
    # TEST AUTOMATION READINESS
    # =============================================================

    def _add_test_readiness(
        self,
        document: Document,
        state: RequirementState,
    ) -> None:
        functional = self._clean(
            state.functional_requirements
        )

        ui_candidates = [
            requirement
            for requirement in functional
            if self._is_ui_candidate(requirement)
        ]

        rows = [
            [
                "Functional Requirements",
                str(len(functional)),
            ],
            [
                "Potential UI Requirements",
                str(len(ui_candidates)),
            ],
            [
                "Potential Playwright Coverage",
                (
                    "Required"
                    if ui_candidates
                    else "Not currently identified"
                ),
            ],
            [
                "Automation Readiness",
                (
                    "Review Required"
                    if state.open_questions
                    else "Candidate for test generation"
                ),
            ],
        ]

        self._table(
            document,
            ["Automation Attribute", "Assessment"],
            rows,
        )

        if ui_candidates:
            self._paragraph(
                document,
                "Potential UI automation requirements:"
            )

            for requirement in ui_candidates:
                self._bullet(
                    document,
                    requirement,
                )

    # =============================================================
    # INFERENCE
    # =============================================================

    @staticmethod
    def _infer_priority(requirement: str) -> str:
        text = requirement.lower()

        if any(
            keyword in text
            for keyword in (
                "must",
                "mandatory",
                "critical",
                "required",
                "shall",
            )
        ):
            return "High"

        if any(
            keyword in text
            for keyword in (
                "should",
                "important",
            )
        ):
            return "Medium"

        if any(
            keyword in text
            for keyword in (
                "could",
                "may",
                "optional",
            )
        ):
            return "Low"

        return "TBD"

    @staticmethod
    def _infer_nfr_category(
        requirement: str,
    ) -> str:
        text = requirement.lower()

        categories = {
            "Performance": (
                "performance",
                "response time",
                "latency",
                "throughput",
                "load",
                "speed",
            ),
            "Security": (
                "security",
                "authentication",
                "authorization",
                "password",
                "encryption",
                "access control",
            ),
            "Availability": (
                "availability",
                "uptime",
                "downtime",
                "recovery",
            ),
            "Scalability": (
                "scalability",
                "scale",
                "concurrent",
                "users",
            ),
            "Usability": (
                "usability",
                "user friendly",
                "accessible",
                "accessibility",
            ),
            "Reliability": (
                "reliable",
                "reliability",
                "fault tolerance",
                "failure",
            ),
            "Maintainability": (
                "maintainability",
                "maintain",
                "supportability",
            ),
        }

        for category, keywords in categories.items():
            if any(
                keyword in text
                for keyword in keywords
            ):
                return category

        return "General"

    @staticmethod
    def _infer_nfr_target(
        requirement: str,
        category: str,
    ) -> str:
        text = requirement.lower()

        # Look for explicit numeric targets.
        import re

        numeric = re.search(
            r"\b\d+(?:\.\d+)?\s*(?:ms|sec|seconds|minutes|%|gb|mb|users?)\b",
            text,
        )

        if numeric:
            return numeric.group(0)

        defaults = {
            "Performance": (
                "Specific response-time target required."
            ),
            "Security": (
                "Security control must be verified against "
                "the applicable security requirements."
            ),
            "Availability": (
                "Specific availability target required."
            ),
            "Scalability": (
                "Specific concurrency/scalability target required."
            ),
            "Usability": (
                "Usability criteria require definition."
            ),
            "Reliability": (
                "Reliability criteria require definition."
            ),
            "Maintainability": (
                "Maintainability criteria require definition."
            ),
        }

        return defaults.get(
            category,
            "Specific measurable target required.",
        )

    @staticmethod
    def _generate_acceptance_criteria(
        requirement: str,
    ) -> str:
        return (
            "Given the required preconditions, when the specified "
            "behavior is executed, then the expected outcome described "
            "by the requirement shall be achieved."
        )

    @staticmethod
    def _verification_method(
        requirement: str,
    ) -> str:
        text = requirement.lower()

        if any(
            keyword in text
            for keyword in (
                "screen",
                "page",
                "button",
                "form",
                "login",
                "click",
                "upload",
                "download",
                "ui",
                "user interface",
            )
        ):
            return "UI / Playwright"

        if any(
            keyword in text
            for keyword in (
                "api",
                "endpoint",
                "request",
                "response",
            )
        ):
            return "API / Integration Test"

        if any(
            keyword in text
            for keyword in (
                "performance",
                "latency",
                "response time",
            )
        ):
            return "Performance Test"

        if any(
            keyword in text
            for keyword in (
                "security",
                "authentication",
                "authorization",
            )
        ):
            return "Security Test"

        return "Functional Test"

    @staticmethod
    def _is_ui_candidate(
        requirement: str,
    ) -> bool:
        text = requirement.lower()

        keywords = (
            "ui",
            "user interface",
            "screen",
            "page",
            "button",
            "form",
            "field",
            "dropdown",
            "checkbox",
            "login",
            "upload",
            "download",
            "click",
            "select",
            "enter",
            "display",
        )

        return any(
            keyword in text
            for keyword in keywords
        )

    @staticmethod
    def _has_explicit_priority(
        requirement: str,
    ) -> bool:
        text = requirement.lower()

        return any(
            keyword in text
            for keyword in (
                "must",
                "mandatory",
                "critical",
                "required",
                "shall",
                "should",
                "important",
                "optional",
                "could",
                "may",
            )
        )

    # =============================================================
    # TEXT / LIST HELPERS
    # =============================================================

    @staticmethod
    def _clean(
        values: Iterable[str] | None,
    ) -> list[str]:
        if not values:
            return []

        return [
            str(value).strip()
            for value in values
            if value is not None
            and str(value).strip()
        ]

    @classmethod
    def _value(
        cls,
        value: str | None,
    ) -> str:
        if value and str(value).strip():
            return str(value).strip()

        return cls.TBD

    @classmethod
    def _text(
        cls,
        document: Document,
        value: str | None,
        fallback: str | None = None,
    ) -> None:
        if value and str(value).strip():
            document.add_paragraph(
                str(value).strip()
            )
            return

        document.add_paragraph(
            fallback or cls.TBD
        )

    @staticmethod
    def _paragraph(
        document: Document,
        text: str,
    ) -> None:
        document.add_paragraph(text)

    @staticmethod
    def _bullet(
        document: Document,
        text: str,
    ) -> None:
        document.add_paragraph(
            text,
            style="List Bullet",
        )

    @classmethod
    def _list(
        cls,
        document: Document,
        values: list[str] | None,
        fallback: str,
    ) -> None:
        cleaned = cls._clean(values)

        if not cleaned:
            cls._bullet(
                document,
                fallback,
            )
            return

        for value in cleaned:
            cls._bullet(
                document,
                value,
            )

    # =============================================================
    # STAKEHOLDER INFERENCE
    # =============================================================

    @staticmethod
    def _stakeholder_responsibility(
        stakeholder: str,
    ) -> str:
        return (
            "Provide business direction, requirements clarification, "
            "review, and approval."
        )

    @staticmethod
    def _role_responsibility(
        role: str,
    ) -> str:
        role_lower = role.lower()

        if "admin" in role_lower:
            return (
                "Manage administrative functions, configuration, "
                "and access according to defined permissions."
            )

        if "tester" in role_lower or "qa" in role_lower:
            return (
                "Validate requirements through functional and "
                "automated testing."
            )

        if "developer" in role_lower:
            return (
                "Implement and maintain the solution according "
                "to approved requirements."
            )

        if "manager" in role_lower:
            return (
                "Review scope, priorities, progress, and "
                "business acceptance."
            )

        return (
            "Perform activities associated with the defined "
            "user role."
        )

    # =============================================================
    # TABLE
    # =============================================================

    @staticmethod
    def _table(
        document: Document,
        headers: list[str],
        rows: list[list[str]],
    ) -> None:
        table = document.add_table(
            rows=1,
            cols=len(headers),
        )

        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        table.style = "Table Grid"

        # Header
        for cell, header in zip(
            table.rows[0].cells,
            headers,
        ):
            cell.text = header

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            SRDGenerator._shade(
                cell,
                "1F4E5F",
            )

            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(
                    255,
                    255,
                    255,
                )

        # Body
        for row in rows:
            cells = table.add_row().cells

            for cell, value in zip(
                cells,
                row,
            ):
                cell.text = (
                    str(value)
                    if value is not None
                    else SRDGenerator.TBD
                )

                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

        # Borders
        for row in table.rows:
            for cell in row.cells:
                SRDGenerator._borders(cell)

        document.add_paragraph()

    # =============================================================
    # FORMATTING
    # =============================================================

    @staticmethod
    def _heading(
        document: Document,
        text: str,
    ) -> None:
        document.add_heading(
            text,
            level=1,
        )

    @staticmethod
    def _subheading(
        document: Document,
        text: str,
    ) -> None:
        document.add_heading(
            text,
            level=2,
        )

    @staticmethod
    def _configure_footer(
        document: Document,
    ) -> None:
        section = document.sections[0]

        footer = section.footer.paragraphs[0]

        footer.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        footer.add_run(
            "AI Requirements Wizard | "
            "Software Requirements Specification"
        )

    @staticmethod
    def _shade(
        cell,
        fill: str,
    ) -> None:
        properties = cell._tc.get_or_add_tcPr()

        shading = OxmlElement("w:shd")

        shading.set(
            qn("w:fill"),
            fill,
        )

        properties.append(shading)

    @staticmethod
    def _borders(
        cell,
    ) -> None:
        properties = cell._tc.get_or_add_tcPr()

        borders = (
            properties.first_child_found_in(
                "w:tcBorders"
            )
            or OxmlElement("w:tcBorders")
        )

        if borders.getparent() is None:
            properties.append(borders)

        for edge in (
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ):
            element = (
                borders.find(
                    qn("w:" + edge)
                )
                or OxmlElement(
                    "w:" + edge
                )
            )

            if element.getparent() is None:
                borders.append(element)

            element.set(
                qn("w:val"),
                "single",
            )

            element.set(
                qn("w:sz"),
                "4",
            )

            element.set(
                qn("w:color"),
                "D9D9D9",
            )