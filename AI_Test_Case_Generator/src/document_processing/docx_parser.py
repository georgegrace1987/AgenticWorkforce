from pathlib import Path
from typing import List
from docx import Document
from src.models.document_models import DocumentModel, ParagraphModel, TableModel, SourceLocation
from src.utils.logger import get_logger

logger = get_logger("docx_parser")


def parse_docx(path: Path) -> DocumentModel:
    """Parse a DOCX file into a DocumentModel, extracting paragraphs and tables."""
    doc = Document(str(path))
    paragraphs: List[ParagraphModel] = []
    tables: List[TableModel] = []

    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            loc = SourceLocation(filename=str(path), section=f"paragraph_{i}")
            paragraphs.append(ParagraphModel(text=text, location=loc))

    for t_index, table in enumerate(doc.tables, start=1):
        headers = []
        rows = []
        for r_index, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if r_index == 0:
                headers = cells
            else:
                rows.append(cells)
        loc = SourceLocation(filename=str(path), section=f"table_{t_index}")
        tables.append(TableModel(headers=headers or [f"col_{i}" for i in range(len(rows[0]) if rows else 0)], rows=rows, location=loc))

    document = DocumentModel(
        filename=str(path.name),
        document_type="docx",
        paragraphs=paragraphs,
        tables=tables,
        metadata={},
        source_locations=[{"filename": str(path)}],
    )
    return document
